import copy
import docx

from ..html import get_html_tree
from ..word_docx.tables import replace_in_table
from ..word_docx.links import create_hyperlink

def get_work_program(Q, verbose: bool = False, studyGroup: int = 12, isn_sp: int|None = None):
    isn_sp = False

    if isn_sp:
        # Preparing for the day the ITU-T API is being updated to only allow ISNs
        sgIndex = {12: 9683}
        qIndex = {
            1: 10694,
            2: 10695,
            4: 10696,
            5: 10697,
            6: 10698,
            7: 10699,
            9: 10700,
            10: 10701,
            12: 10702,
            13: 10703,
            14: 10704,
            15: 10705,
            17: 10706,
            19: 10707,
            20: 10708,
        }

        url = f"https://www.itu.int/ITU-T/workprog/wp_search.aspx?isn_sg={sgIndex[studyGroup]}&qIndex={qIndex[Q]}&isn_sp={isn_sp}&isn_status=-1,1,3,7&details=1&view=tab&field=ahjgoflki"
    else:
        url = f"https://www.itu.int/ITU-T/workprog/wp_search.aspx?sg={studyGroup}&q={Q}&isn_sp={isn_sp}&isn_status=-1,1,3,7&details=1&view=tab&field=ahjgoflki"

    info = []

    if verbose:
        print(f"  Fetching work program from: {url}")
    tree = get_html_tree(url)

    # Find and parse all rows (<tr>) in the document
    rows = tree.xpath("//table[contains(@id, 'tab_tabular_view_gd_wp_tabular')]/tr")

    if not len(rows) >= 1:
        print("Error fetching work program - Please check query response")
        raise ()

    for row in rows:
        item = {}

        try:
            tds = row.xpath(".//td")

            # 1st column - Work item
            item["href"] = tds[0].xpath(".//a/@href")[0].strip()
            item["work_item"] = tds[0].xpath(".//a/text()")[0]

            # 2nd column - Version
            item["version"] = tds[1].xpath(".//div/text()")[0]

            # 3rd column - Title
            item["title"] = tds[2].xpath(".//text()")[0]

            # 4th column - Approval process
            item["process"] = tds[3].xpath(".//div/text()")[0]

            # 5th column -  Priority
            item["priority"] = tds[4].xpath(".//div/text()")[0]

            # 6th column -  Timing
            item["timing"] = tds[5].xpath(".//div/nobr/text()")[0]

            # 7th column -  Editors
            try:
                item["editors"] = []
                for editor in tds[6].xpath(".//a"):
                    tmp = {}
                    tmp["href"] = editor.xpath(".//@href")[0].strip().replace("(AT)", "@")
                    tmp["name"] = editor.xpath("./text()")[0]
                    item["editors"].append(tmp)
            except:
                print("!!! Cannot retrieve editors")
                pass

            # 8th column - base document(s)
            texts = tds[8].xpath(".//a")

            item["basetext"] = []
            for text in tds[7].xpath(".//a"):
                tmp = {}
                tmp["href"] = text.xpath(".//@href")[0].strip()
                tmp["name"] = text.xpath(".//text()")[0]
                item["basetext"].append(tmp)

            # 9th column - Liaisons
            item["relationship"] = [x.strip() for x in tds[8].xpath(".//text()")[0].split(",")]

            info.append(item)

        except Exception as e:
            pass

    return info

def insert_work_program(document, info):
    # Fid the work program table
    targetTable = None

    for table in document.tables:
        for idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if targetTable != None:
                        break
                    if paragraph.text == "Approval process":
                        targetTable = table

    for idx, work_item in enumerate(info):
        # Duplicate row
        if idx != len(info) - 1:
            targetTable.rows[-1]._tr.addnext(copy.deepcopy(targetTable.rows[-1]._tr))

        new_h = create_hyperlink(targetTable, work_item["work_item"], work_item["href"], format="hyperlink")
        replace_in_table(targetTable, "WP_WorkItem", new_h)
        replace_in_table(targetTable, "WP_Version", work_item["version"])
        replace_in_table(targetTable, "WP_Process", work_item["process"])
        replace_in_table(targetTable, "WP_Priority", work_item["priority"])
        replace_in_table(targetTable, "WP_Timing", work_item["timing"])
        replace_in_table(targetTable, "WP_Relationship", ",\n".join([x for x in work_item["relationship"]]))
        replace_in_table(targetTable, "WP_Title", work_item["title"])
        replace_in_table(targetTable, "WP_Editors", ",\n".join([x["name"] for x in work_item["editors"]]))

        # Generate base texts with links
        newRun = docx.oxml.shared.OxmlElement("w:r")
        for idx, x in enumerate(work_item["basetext"]):
            if idx > 0:
                tmp = docx.oxml.shared.OxmlElement("w:r")
                tmp.text = ", "
                newRun.append(tmp)
            newRun.append(create_hyperlink(document, x["name"], x["href"], format="hyperlink"))

        replace_in_table(targetTable, "WP_BaseTexts", newRun)

    pass

if __name__ == "__main__":
    pass