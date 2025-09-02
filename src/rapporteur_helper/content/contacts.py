import copy
from typing import Any

from docx.document import Document

from ..word_docx.paragraph import replace
from ..word_docx.tables import replace_in_table


def insert_contacts(document: Document, questionInfo: dict[str, Any]):
    numContacts = len(questionInfo["rapporteurs"])

    # Fid the contact table
    contactTable = None
    for table in document.tables:
        for idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if contactTable != None:
                        break
                    if paragraph.text == "Contact:":
                        contactTable = table

    # Add contacts row if necessary (there are two in the template)
    for i in range(0, numContacts - 2, 1):
        contactTable.rows[-1]._tr.addnext(copy.deepcopy(contactTable.rows[-1]._tr))

    if numContacts == 1:
        contactTable._tbl.remove(contactTable.rows[-1]._tr)

    # Update the contact table
    for contact in questionInfo["rapporteurs"]:
        replace_in_table(contactTable, "Name", f"{contact['firstName']} {contact['lastName']}")
        replace_in_table(contactTable, "Organization", f"{contact['company']}")
        replace_in_table(contactTable, "Country", f"{contact['country']}")
        if "tel" in contact:
            replace_in_table(contactTable, "Tel:\t+xx", f"Tel:\t{contact['tel']}")
        else:
            replace_in_table(contactTable, "Tel:\t+xx", "")
        replace_in_table(contactTable, "a@b.com", f"{contact['email']}")


def _get_contact_line(contact: dict) -> str:
    return f"{contact['firstName']} {contact['lastName']} ({contact['company']}, {contact['country']})"


def _concat_contact_lines(contact_lines: list[dict | str]) -> str:
    # Concatenate contact lines with appropriate conjunctions
    if not contact_lines:
        return ""

    lines = [_get_contact_line(c) for c in contact_lines] if all(isinstance(c, dict) for c in contact_lines) else [str(c) for c in contact_lines]

    if len(lines) == 1:
        return lines[0]
    return ", ".join(lines[:-1]) + " and " + lines[-1]


def get_chair_text(rapporteur_info: list[dict[str, str]]) -> str:
    # Format text for Section 1:

    # check: associate(s) vs co-rapporteur(s)?
    rapporteurs = []
    associates = []
    for contact in rapporteur_info:
        if "associate" in contact["role"].lower():
            associates.append(contact)
        else:
            rapporteurs.append(contact)

    text = "rapporteur" if len(rapporteurs) == 1 else "co-rapporteurs"
    text += f" {_concat_contact_lines(rapporteurs)}"

    if len(associates) > 0:
        text += f", with the assistance of {_concat_contact_lines(associates)}"

    return text


if __name__ == "__main__":
    pass
