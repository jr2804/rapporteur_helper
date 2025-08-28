import requests
import re
from typing import Any
from lxml import html
from docx.text.paragraph import Paragraph

from . import hostname
from ..word_docx.links import add_hyperlink

def insert_documents(docSection: Paragraph, endpoints: list|Any, verbose: bool = False, studyGroup: int = 12):
    if not isinstance(endpoints, list):
        endpoints = [endpoints]

    rows = []
    for endpoint in endpoints:
        if verbose:
            print(f"  Retrieving documents from: {endpoint['url']}")
        x = requests.get(endpoint["url"])
        tree = html.fromstring(x.content)

        # Find and parse all rows (<tr>) in the document
        rows += tree.xpath("//tr")

    docSection.text = ""

    # Parse  in descending order
    for i in range(len(rows) - 1, 0, -1):
        row = rows[i]
        columns = row.xpath(".//td")

        try:
            # A document row should have the attributes below
            # If not, then the row is ignored

            # Link and document number should be in the second column
            link = hostname + "/" + columns[1].xpath(".//a")[0].attrib["href"].strip()

            try:
                number = columns[1].xpath(".//a/strong/text()")[0].strip().replace("[ ", endpoint["prefix"]).replace(" ]", "")
            except:
                # Handle case where documents have not yet been uploaded
                number = columns[1].xpath(".//a/text()")[0].strip().replace("[ ", endpoint["prefix"]).replace(" ]", "")

            try:
                revision = columns[1].xpath(".//font/text()")[0]
                x = re.search(r"([\d]+)\)", revision)
                revision = x.group(1)
                number = f"{number}r{revision}"
            except Exception as e:
                # print(e)
                pass

            # Title should be in third row
            title = columns[2].xpath(".//text()")[0].strip()
            sources = columns[3].xpath(".//a")
            src = []
            for source in sources:
                src.append(dict(link=f"{hostname}/{source.attrib['href']}", text=source.text.strip()))

            # Relevant questions should be in fourth column
            questions = columns[4].xpath(".//a")
            q = []
            for quest in questions:
                q.append(dict(link=f"{hostname}/{quest.attrib['href']}", text=quest.text.strip().replace(f"/{studyGroup}", "")))

            # Generate word document block for this document
            # p = document.add_paragraph()
            p = docSection.insert_paragraph_before()

            tmpNumber = number.replace("-GEN", "")
            add_hyperlink(p, f"{tmpNumber} - {title}", link, format="bold")

            p.add_run("\nSources: ")
            for item in src:
                add_hyperlink(p, item["text"], item["link"], format="hyperlink")
                if src[-1] != item:
                    p.add_run(" | ")

            p.add_run("\nQuestions: ")
            for item in q:
                add_hyperlink(p, item["text"], item["link"], format="hyperlink")
                if q[-1] != item:
                    p.add_run(", ")

            # Do not include a summary section for documents addressed to Q.ALL
            if q[0]["link"].find("QALL") < 0:
                p.add_run("\nSummary:\n")

            p.add_run("\n")

        except Exception as e:
            # print(e)
            # traceback.print_exc()
            pass


if __name__ == "__main__":
    pass