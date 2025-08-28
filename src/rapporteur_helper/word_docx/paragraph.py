
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph

def insert_paragraph_after(paragraph: Paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = docx.oxml.shared.OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph(new_p, paragraph)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def find_element(document: Document, text) -> Paragraph|None:
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph


def replace(document: Document, find: str, replace: str):
    for paragraph in document.paragraphs:
        foundInRun = False
        for run in paragraph.runs:
            if find in run.text:
                run.text = run.text.replace(find, replace)
                # print(f'run: {find}')
                foundInRun = True
        if foundInRun == False:
            if find in paragraph.text:
                paragraph.text = paragraph.text.replace(find, replace)
                # print(f'paragraph: {find}')

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    foundInRun = False
                    for run in paragraph.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
                            # print(f'run: {find}')
                            foundInRun = True
                    if foundInRun == False:
                        if find in paragraph.text:
                            paragraph.text = paragraph.text.replace(find, replace)
                            # print(f'paragraph: {find}')

if __name__ == "__main__":
    pass