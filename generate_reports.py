from docx import Document
import docx
import copy

import os
from lxml import html, etree
import requests
from pprint import pprint
import re
import traceback

verbose = True

questions = range(1, 21)
# questions = [1,2, 7, 14]
# meetingDate = "220607"
add_qall = False

# Update these parameters for each meeting
studyGroup = 12
meetingDetails = "Geneva, 14-23 January 2025"
meetingDate = "250114"

# Update these parameters to the current study period 
studyPeriodId = 18
studyPeriodStart = 25
isn_sp = 9677

def create_hyperlink(document, text, url, format = ['None', 'bold', 'italic', 'hyperlink', 'button'][0]):
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # Access to the document settings (DocumentPart) to create a new relation id value
    documentPart = document.part
    r_id = documentPart.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Attach the relation ID to the hyperlink object
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    if format == 'italic':
        rStyle = docx.oxml.shared.OxmlElement('w:i')
        rPr.append(rStyle)
    if format == 'bold':
        rStyle = docx.oxml.shared.OxmlElement('w:b')
        rPr.append(rStyle)

    if format == 'hyperlink':
        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

    # Join all the xml elements together and add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = f"{text}"

    hyperlink.append(new_run)

    if format == 'button':
        # Create a link button with the Hyperlink style, prettier than coloring the whole text
        new_run = docx.oxml.shared.OxmlElement('w:r')
        new_run.text = "  "
        hyperlink.append(new_run)

        new_run = docx.oxml.shared.OxmlElement('w:r')
        # Unicode character for "download"
        # https://www.fileformat.info/info/unicode/char/2913/fontsupport.htm

        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rPr.append(rStyle)
        new_run.append(rPr)

        new_run.text = '\u2913'

        hyperlink.append(new_run)

    return hyperlink

def add_hyperlink(paragraph, text, url, format = ['None', 'bold', 'italic', 'hyperlink', 'button'][0]):
    # :param paragraph: The paragraph we are adding the hyperlink to.
    # :param text: The text displayed for the url
    # :param url: A string containing the required url
    # :param format: Style to apply to the text ['None', 'bold','italic']
    #     :return: The hyperlink object

    # Create hyperlink
    hyperlink = create_hyperlink(paragraph, text, url, format)

    # Append hyperlink to paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = docx.oxml.shared.OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph(new_p, paragraph)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def insert_documents(docSection, endpoints):
    if not isinstance(endpoints, list):
        endpoints = [endpoints]

    rows = []
    for endpoint in endpoints:
        if verbose:
            print(f"  Retrieving documents from: {endpoint['url']}")
        x = requests.get(endpoint['url'])
        tree = html.fromstring(x.content)

        # Find and parse all rows (<tr>) in the document
        rows += tree.xpath('//tr')

    docSection.text = ""

    # Parse  in descending order
    for i in range(len(rows) -1, 0, -1):
        row = rows[i]
        columns = row.xpath('.//td')

        try:
            # A document row should have the attributes below
            # If not, then the row is ignored

            # Link and document number should be in the second column
            link = hostname + '/' + columns[1].xpath('.//a')[0].attrib['href'].strip()

            try:
                number = columns[1].xpath('.//a/strong/text()')[0].strip().replace('[ ', endpoint['prefix']).replace(' ]', '')
            except:
                # Handle case where documents have not yet been uploaded
                number = columns[1].xpath('.//a/text()')[0].strip().replace('[ ', endpoint['prefix']).replace(' ]', '')

            try:
                revision = columns[1].xpath('.//font/text()')[0]
                x = re.search(r"([\d]+)\)", revision)
                revision = x.group(1)
                number = f"{number}r{revision}"
            except Exception as e:
                # print(e)
                pass

            # Title should be in third row
            title = columns[2].xpath('.//text()')[0].strip()
            sources = columns[3].xpath('.//a')
            src = []
            for source in sources:
                src.append(dict(link = f'{hostname}/{source.attrib["href"]}', text = source.text.strip()))

            # Relevant questions should be in fourth column
            questions = columns[4].xpath('.//a')
            q = []
            for quest in questions:
                q.append(dict(link = f'{hostname}/{quest.attrib["href"]}', text = quest.text.strip().replace(f'/{studyGroup}', '')))

            # Generate word document block for this document
            # p = document.add_paragraph()
            p = docSection.insert_paragraph_before()

            tmpNumber = number.replace('-GEN', '')
            add_hyperlink(p, f"{tmpNumber} - {title}", link, format = 'bold')

            p.add_run('\nSources: ')
            for item in src:
                add_hyperlink(p, item['text'], item['link'], format='hyperlink')
                if src[-1] != item:
                    p.add_run(' | ')

            p.add_run('\nQuestions: ')
            for item in q:
                add_hyperlink(p, item['text'], item['link'], format='hyperlink')
                if q[-1] != item:
                    p.add_run(', ')

            # Do not include a summary section for documents addressed to Q.ALL
            if q[0]['link'].find('QALL') < 0:
                p.add_run('\nSummary:\n')

            p.add_run('\n')

        except Exception as e:
            # print(e)
            # traceback.print_exc()
            pass

def get_html_tree(url):
    try:
        x = requests.get(url)
        tree = html.fromstring(x.content)
        return tree
    except Exception as e:
        print(url)
        raise(e)

def get_work_program(Q):
    useISN = False

    if useISN:
        # Preparing for the day the ITU-T API is being updated to only allow ISNs
        sgIndex = {12: 9683}
        qIndex = { 1: 10694, 2: 10695, 4: 10696, 5: 10697, 6: 10698, 7: 10699, 9: 10700,
                   10: 10701, 12: 10702, 13: 10703, 14: 10704, 15: 10705, 17: 10706,
                   19: 10707, 20: 10708 }

        url = f"https://www.itu.int/ITU-T/workprog/wp_search.aspx?isn_sg={sgIndex[studyGroup]}&qIndex={qIndex[Q]}&isn_sp={isn_sp}&isn_status=-1,1,3,7&details=1&view=tab&field=ahjgoflki"
    else:
        url = f"https://www.itu.int/ITU-T/workprog/wp_search.aspx?sg={studyGroup}&q={Q}&isn_sp={isn_sp}&isn_status=-1,1,3,7&details=1&view=tab&field=ahjgoflki"

    info = []

    if verbose:
        print(f'  Fetching work program from: {url}')
    tree = get_html_tree(url)

    # Find and parse all rows (<tr>) in the document
    rows = tree.xpath("//table[contains(@id, 'tab_tabular_view_gd_wp_tabular')]/tr")

    if not len(rows) >= 1:
        print("Error fetching work program - Please check query response")
        raise()

    for row in rows:
        item = {}

        try:
            tds = row.xpath(".//td")

            # 1st column - Work item
            item['href'] = tds[0].xpath(".//a/@href")[0].strip()
            item['work_item'] = tds[0].xpath(".//a/text()")[0]

            # 2nd column - Version
            item['version'] = tds[1].xpath(".//div/text()")[0]

            # 3rd column - Title
            item['title'] = tds[2].xpath(".//text()")[0]

            # 4th column - Approval process
            item['process'] = tds[3].xpath(".//div/text()")[0]

            # 5th column -  Priority
            item['priority'] = tds[4].xpath(".//div/text()")[0]

            # 6th column -  Timing
            item['timing'] = tds[5].xpath(".//div/nobr/text()")[0]

            # 7th column -  Editors
            try:
                item['editors'] = []
                for editor in tds[6].xpath(".//a"):
                    tmp = {}
                    tmp['href'] = editor.xpath(".//@href")[0].strip().replace('(AT)', '@')
                    tmp['name'] = editor.xpath("./text()")[0]
                    item['editors'].append(tmp)
            except:
                print('!!! Cannot retrieve editors')
                pass

            # 8th column - base document(s)
            texts = tds[8].xpath(".//a")

            item['basetext'] = []
            for text in tds[7].xpath(".//a"):
                tmp = {}
                tmp['href'] = text.xpath(".//@href")[0].strip()
                tmp['name'] = text.xpath(".//text()")[0]
                item['basetext'].append(tmp)

            # 9th column - Liaisons
            item['relationship'] = [ x.strip() for x in tds[8].xpath(".//text()")[0].split(',')]

            info.append(item)

        except Exception as e:
            pass

    return info

def get_questions_details():
    info = {}
    url=f"https://www.itu.int/net4/ITU-T/lists/loqr.aspx?Group={studyGroup}&Period={studyPeriodId}"

    tree = get_html_tree(url)

    # Find and parse all rows (<tr>) in the document
    rows = tree.xpath('//tr')

    currentQuestion = -1
    for row in rows:
        # Extract WP number and Question title
        try:
            # Question number and WP
            tmp = row.xpath(".//span[contains(@id,'lblQWP')]/text()")[0]

            try:
                res = re.search(fr'Q(\d+)/{studyGroup}.*WP(\d+)/{studyGroup}', tmp)
                qNum = int(res.group(1))
                wpNum = int(res.group(2))
            except:
                # If it fails, check that it is because there is no WP number
                res = re.search(fr'Q(\d+)/{studyGroup}.*', tmp)
                qNum = int(res.group(1))
                wpNum = -1

            # Question title
            qTitle = row.xpath(".//span[contains(@id,'lblQuestion')]/text()")[0]

            if qNum not in info:
                info[qNum] = dict(rapporteurs = [])

            info[qNum].update(dict(wp = wpNum, title = qTitle))

            currentQuestion = qNum

            # print(info[qNum])
        except Exception as e:
            # print(e)
            pass

        # Extract Rapporteurs contact details
        try:
            tmp = {}
            tmp['firstName'] = row.xpath(".//span[contains(@id,'dtlRappQues_lblFName')]/text()")[0]
            tmp['lastName'] = row.xpath(".//span[contains(@id,'dtlRappQues_lblLName')]/text()")[0].upper()
            tmp['role'] = row.xpath(".//span[contains(@id,'dtlRappQues_lblRole')]/text()")[0]
            tmp['company'] = row.xpath(".//span[contains(@id,'dtlRappQues_lblCompany')]/text()")[0]
            tmp['address'] = ' '.join(row.xpath(".//span[contains(@id,'dtlRappQues_lblAddress')]/text()"))
            tmp['country'] = row.xpath(".//span[contains(@id,'dtlRappQues_lblAddress')]/text()")[-1]
            
            try:
                # Some Rapporteurs do not have a telephone number available
                tmp['tel'] = row.xpath(".//span[contains(@id,'dtlRappQues_telLabel')]/text()")[0]
            except:
                pass
            tmp['email'] = row.xpath(".//a[contains(@id,'dtlRappQues_linkemail')]/text()")[0].replace('[at]', '@')
            
            info[currentQuestion]['rapporteurs'].append(tmp)
        except Exception as e:
            # traceback.print_exc()
            # print(e)
            pass
    
    if len(info) < 1:
        raise(Exception(f'get_questions_details - Could not parse question details from {url}'))

    return info

def find_element(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph

def replace(find, replace):
    for paragraph in document.paragraphs:
        foundInRun = False
        for run in paragraph.runs:
            if find in run.text:
                run.text = run.text.replace(find, replace)
                # print(f'run: {find}')
                foundInRun = True
        if foundInRun == False:
            if find in paragraph.text:
                paragraph.text = paragraph.text.replace(find, replace)
                # print(f'paragraph: {find}')

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    foundInRun = False
                    for run in paragraph.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
                            # print(f'run: {find}')
                            foundInRun = True
                    if foundInRun == False:
                        if find in paragraph.text:
                            paragraph.text = paragraph.text.replace(find, replace)
                            # print(f'paragraph: {find}')


def replace_in_table(table, find, replace):
    for row in table.rows:
        for cell in row.cells:
            for subtable in cell.tables:
                if replace_in_table(subtable, find, replace):
                    return True
            for paragraph in cell.paragraphs:
                foundInRun = False
                for run in paragraph.runs:
                    if find in run.text:
                        if isinstance(replace, str):
                            run.text = run.text.replace(find, replace)
                            run.font.highlight_color = 0
                        else:
                            run.text = ''
                            paragraph._p.append(replace)
                        return True

                if foundInRun == False:
                    if find in paragraph.text:
                        if isinstance(replace, str):
                            paragraph.text = paragraph.text.replace(find, replace)
                        else:
                            pass
                            paragraph._p.addnext(replace)

                        return True
    return False

def insert_contacts(document, questionInfo):
    numContacts = len(questionInfo['rapporteurs'])

    # Fid the contact table
    contactTable = None
    for table in document.tables:
        for idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if contactTable != None:
                        break
                    if paragraph.text == "Contact:":
                        contactTable = table

    # Add contacts row if necessary (there are two in the template)
    for i in range(0, numContacts - 2, 1):
        contactTable.rows[-1]._tr.addnext(copy.deepcopy(contactTable.rows[-1]._tr))

    if numContacts == 1:
        contactTable._tbl.remove(contactTable.rows[-1]._tr)

    # Update the contact table
    for contact in questionInfo['rapporteurs']:
        replace_in_table(contactTable, "Name", f"{contact['firstName']} {contact['lastName']}")
        replace_in_table(contactTable, "Organization", f"{contact['company']}")
        replace_in_table(contactTable, "Country", f"{contact['country']}")
        if 'tel' in contact:
            replace_in_table(contactTable, "Tel:\t+xx", f"Tel:\t{contact['tel']}")
        else:
            replace_in_table(contactTable, "Tel:\t+xx", "")
        replace_in_table(contactTable, "a@b.com", f"{contact['email']}")

    # Format text for Section 1:
    target = "the [co-] chairmanship of name of Rapporteur (organization, country) [with the assistance of name of associate Rapporteur (organization, country)]"

    if numContacts == 1:
        text = f"the chairmanship of {contact['firstName']} {contact['lastName']} ({contact['company']}, {contact['country']})"
    else:
        hasAssociate = False
        for contact in questionInfo['rapporteurs']:
            if "Associate" in contact['role']:
                hasAssociate = True

        if hasAssociate == False:
            text = "the co-chairmanship of "
            tmp = []
            for contact in questionInfo['rapporteurs']:
                tmp.append(f"{contact['firstName']} {contact['lastName']} ({contact['company']}, {contact['country']})")
            text += " and ".join(tmp)
        else:
            text = "the chairmanship of "
            for contact in questionInfo['rapporteurs']:
                if "Rapporteur" in contact['role']:
                    text += f"{contact['firstName']} {contact['lastName']} ({contact['company']}, {contact['country']})"

            for contact in questionInfo['rapporteurs']:
                if "Associate" in contact['role']:
                    text += f" with the assistance of {contact['firstName']} {contact['lastName']} ({contact['company']}, {contact['country']})"

    replace(target, text)

def insert_work_program(document, info):
    # Fid the work program table
    targetTable = None

    for table in document.tables:
        for idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if targetTable != None:
                        break
                    if paragraph.text == "Approval process":
                        targetTable = table

    for (idx, work_item) in enumerate(info):
        # Duplicate row
        if idx != len(info) -1:
            targetTable.rows[-1]._tr.addnext(copy.deepcopy(targetTable.rows[-1]._tr))

        new_h = create_hyperlink(targetTable, work_item['work_item'], work_item['href'], format='hyperlink')
        replace_in_table(targetTable, 'WP_WorkItem', new_h)
        replace_in_table(targetTable, 'WP_Version', work_item['version'])
        replace_in_table(targetTable, 'WP_Process', work_item['process'])
        replace_in_table(targetTable, 'WP_Priority', work_item['priority'])
        replace_in_table(targetTable, 'WP_Timing', work_item['timing'])
        replace_in_table(targetTable, 'WP_Relationship', ",\n".join([ x for x in work_item['relationship']]))
        replace_in_table(targetTable, 'WP_Title', work_item['title'])
        replace_in_table(targetTable, 'WP_Editors', ",\n".join([ x['name'] for x in work_item['editors']]))

        # Generate base texts with links
        newRun = docx.oxml.shared.OxmlElement('w:r')
        for idx,x in enumerate(work_item['basetext']):
            if idx > 0:
                tmp = docx.oxml.shared.OxmlElement('w:r')
                tmp.text = ", "
                newRun.append(tmp)
            newRun.append(create_hyperlink(document, x['name'],  x['href'], format='hyperlink'))

        replace_in_table(targetTable, 'WP_BaseTexts', newRun)

    pass

if __name__ == '__main__':
    try:
        questionInfo = get_questions_details()
        # pprint(questionInfo)
    except Exception as e:
        print("Error - Cannot fetch question details from ITU-T website")
        raise(e)

    for question in questions:
        print(f"\n### Generating report for Q{question}")

        try:
            hostname = 'https://www.itu.int'
            endpoints_c = []
            if add_qall:
                endpoints_c += [dict(url = f'{hostname}/md/meetingdoc.asp?lang=en&parent=T{studyPeriodStart}-SG{studyGroup}-{meetingDate}-C&question=QALL/{studyGroup}', prefix=f'SG{studyGroup}-C', title='Contributions')]
            endpoints_c += [dict(url = f'{hostname}/md/meetingdoc.asp?lang=en&parent=T{studyPeriodStart}-SG{studyGroup}-{meetingDate}-C&question=Q{question}/{studyGroup}', prefix=f'SG{studyGroup}-C', title='Contributions'),]

            endpoints_td = []
            if add_qall:
                endpoints_td += [dict(url = f'{hostname}/md/meetingdoc.asp?lang=en&parent=T{studyPeriodStart}-SG{studyGroup}-{meetingDate}-TD&question=QALL/{studyGroup}', prefix=f'SG{studyGroup}-TD', title='Temporary Documents'),]
            endpoints_td += [dict(url = f'{hostname}/md/meetingdoc.asp?lang=en&parent=T{studyPeriodStart}-SG{studyGroup}-{meetingDate}-TD&question=Q{question}/{studyGroup}', prefix=f'SG{studyGroup}-TD', title='Temporary Documents'),]
            # pprint(endpoints)

            with open('template.docx', 'rb') as f:
                document = Document(f)

            # for style in document.styles:
            #     print(f"{style.name} {style.type}")

            # Meeting date
            replace('[place, dates]', meetingDetails)

            # Abstract
            abstract = f"This document contains the Status report of Question {question}/{studyGroup}: {questionInfo[question]['title']} for the meeting in {meetingDetails}."
            replace('[Insert an abstract]', abstract)

            # Insert contributions
            docSection = find_element(document, 'Copy table of contributions')
            insert_documents(docSection, endpoints_c)

            # Insert temporary documents
            # print("  Inserting temporary documents")
            docSection = find_element(document, 'Copy the TD table')
            insert_documents(docSection, endpoints_td)

            # Replace question number
            replace(f'X/{studyGroup}', f'{question}/{studyGroup}')
            replace(f't{studyPeriodStart}sg{studyGroup}qX@lists.itu.int', f't{studyPeriodStart}sg{studyGroup}q{question}@lists.itu.int')

            # Replace working party number
            replace(f'Working Party y/{studyGroup}', f"Working Party {questionInfo[question]['wp']}/{studyGroup}")

            # Replace question title
            replace('[title of question]', questionInfo[question]['title'])
            replace('Title of question', questionInfo[question]['title'])

            # Insert contacts
            insert_contacts(document, questionInfo[question])

            # Insert work programme
            workProgram = get_work_program(question)
            # pprint(workProgram)
            insert_work_program(document, workProgram)

            try:
                os.mkdir(f'./{meetingDate}')
            except:
                pass

            document.save(f'./{meetingDate}/Q{question}_status_report.docx')

        except Exception as e:
            pprint(e)
            traceback.print_stack()
            # pprint(questionInfo)
